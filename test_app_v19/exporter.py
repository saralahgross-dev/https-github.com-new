from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import tempfile
from generator import heb_num

HEBREW_FONT='Arial'

def set_rtl(p):
    pPr = p._p.get_or_add_pPr()
    bidi = pPr.find(qn('w:bidi'))
    if bidi is None:
        bidi = OxmlElement('w:bidi'); pPr.append(bidi)
    bidi.set(qn('w:val'), '1')
    # Match Word's two Hebrew toolbar selections: RTL paragraph + Align Right.
    jc = pPr.find(qn('w:jc'))
    if jc is None:
        jc = OxmlElement('w:jc'); pPr.append(jc)
    jc.set(qn('w:val'), 'right')

def set_run_rtl(run):
    rPr = run._r.get_or_add_rPr()
    rtl = rPr.find(qn('w:rtl'))
    if rtl is None:
        rtl = OxmlElement('w:rtl'); rPr.append(rtl)
    rtl.set(qn('w:val'),'1')
    lang = rPr.find(qn('w:lang'))
    if lang is None:
        lang = OxmlElement('w:lang'); rPr.append(lang)
    lang.set(qn('w:bidi'), 'he-IL')
    lang.set(qn('w:val'), 'he-IL')

def para(doc, text='', bold=False, size=None, rtl=True, align=None, italic=False, space_after=3):
    p = doc.add_paragraph()
    if rtl: set_rtl(p)
    p.alignment = align if align is not None else (WD_ALIGN_PARAGRAPH.RIGHT if rtl else WD_ALIGN_PARAGRAPH.LEFT)
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text); r.bold=bold; r.italic=italic; r.font.name=HEBREW_FONT
    if rtl: set_run_rtl(r)
    if size: r.font.size=Pt(size)
    return p

def setup_doc():
    doc=Document(); sec=doc.sections[0]
    sec.top_margin=Inches(.55); sec.bottom_margin=Inches(.55); sec.left_margin=Inches(.65); sec.right_margin=Inches(.65)
    styles=doc.styles
    styles['Normal'].font.name=HEBREW_FONT; styles['Normal'].font.size=Pt(11)
    return doc

def marker(i): return heb_num(i)+'.'

def _remove_table_borders(table):
    tblPr = table._tbl.tblPr
    borders = tblPr.first_child_found_in('w:tblBorders')
    if borders is None:
        borders = OxmlElement('w:tblBorders')
        tblPr.append(borders)
    for edge in ('top','left','bottom','right','insideH','insideV'):
        tag = 'w:' + edge
        el = borders.find(qn(tag))
        if el is None:
            el = OxmlElement(tag); borders.append(el)
        el.set(qn('w:val'),'nil')

def name_line(doc):
    """Physically pin the blank line on the left and שם on the far right.
    Using two cells avoids Word bidi reordering the underline and Hebrew label.
    """
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    _remove_table_borders(table)
    left, right = table.rows[0].cells
    left.width = Inches(5.7); right.width = Inches(.65)
    lp = left.paragraphs[0]; lp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    lr = lp.add_run('______________________________'); lr.font.name = HEBREW_FONT; lr.font.size = Pt(11)
    rp = right.paragraphs[0]; set_rtl(rp); rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    rr = rp.add_run('שם'); set_run_rtl(rr); rr.font.name = HEBREW_FONT; rr.font.size = Pt(11)
    return table

def _header(doc, config, title):
    # Match teacher samples: name line at top right, then class, then centered Chumash title.
    name_line(doc)
    class_name=(config.get('class_name') or 'כתה ט׳').strip()
    if class_name: para(doc, class_name, False, 11, True, WD_ALIGN_PARAGRAPH.RIGHT, space_after=2)
    perek=heb_num(int(config.get('perek',1)))
    shown_title=(title or '').strip()
    if not shown_title or shown_title.lower().startswith('shemot') or shown_title.startswith('מבחן שמות'):
        shown_title=f'מבחן בחומש – פרק {perek}'
    para(doc, shown_title, True, 16, True, WD_ALIGN_PARAGRAPH.CENTER, space_after=8)

def _group(test):
    grouped=[]; seen={}
    for q in test:
        sec=q.get('section','שאלות')
        if sec not in seen:
            seen[sec]=[]; grouped.append((sec,seen[sec]))
        seen[sec].append(q)
    return grouped

def export_test_docx(test, config, title):
    fd,path=tempfile.mkstemp(suffix='.docx'); Path(path).unlink(missing_ok=True)
    doc=setup_doc(); bilingual=config.get('test_version')=='ivrit_english'; _header(doc,config,title)
    section_no=1
    for section,qs in _group(test):
        para(doc, f'{heb_num(section_no)}. {section}', True, 13, True, WD_ALIGN_PARAGRAPH.RIGHT, space_after=4)
        for i,q in enumerate(qs,1):
            source_tag=''
            if q.get('perek') and q.get('pasuk'):
                source_tag=f" ({heb_num(int(q['perek']))}, {heb_num(int(q['pasuk']))})"
            para(doc, f'{marker(i)} {q.get("prompt","")}{source_tag}', False, 11, True, WD_ALIGN_PARAGRAPH.RIGHT, space_after=1)
            if bilingual and q.get('prompt_en'):
                para(doc, q.get('prompt_en',''), False, 10, False, WD_ALIGN_PARAGRAPH.LEFT, italic=False, space_after=2)
            para(doc, '________________________________________________________________________', False, 10, True, WD_ALIGN_PARAGRAPH.RIGHT, space_after=0)
            if section in ('שאלות קצרות','מפרשים','שאלת ותשובת רש״י'):
                para(doc, '________________________________________________________________________', False, 10, True, WD_ALIGN_PARAGRAPH.RIGHT, space_after=3)
        section_no+=1
    para(doc,'בהצלחה רבה !!!',True,11,True,WD_ALIGN_PARAGRAPH.CENTER,space_after=2)
    doc.save(path); return Path(path)

def export_answer_key_docx(test, config, title):
    fd,path=tempfile.mkstemp(suffix='.docx'); Path(path).unlink(missing_ok=True)
    doc=setup_doc(); bilingual=config.get('test_version')=='ivrit_english'; _header(doc,config,title)
    para(doc,'מפתח תשובות',True,15,True,WD_ALIGN_PARAGRAPH.CENTER,space_after=8)
    section_no=1
    for section,qs in _group(test):
        para(doc,f'{heb_num(section_no)}. {section}',True,13,True,WD_ALIGN_PARAGRAPH.RIGHT,space_after=4)
        for i,q in enumerate(qs,1):
            para(doc,f'{marker(i)} {q.get("prompt","")}',True,10,True,WD_ALIGN_PARAGRAPH.RIGHT,space_after=1)
            if bilingual and q.get('prompt_en'): para(doc,q.get('prompt_en',''),False,9,False,WD_ALIGN_PARAGRAPH.LEFT,space_after=1)
            para(doc,'תשובה: '+q.get('answer',''),False,10,True,WD_ALIGN_PARAGRAPH.RIGHT,space_after=1)
            if bilingual and q.get('answer_en'): para(doc,'Answer: '+q.get('answer_en',''),False,9,False,WD_ALIGN_PARAGRAPH.LEFT,space_after=1)
            if q.get('source_ref'): para(doc,'מקור: '+q.get('source_ref',''),False,9,True,WD_ALIGN_PARAGRAPH.RIGHT,space_after=3)
        section_no+=1
    doc.save(path); return Path(path)

def export_docx(test,config,title=None):
    return export_test_docx(test,config,title)
