import io, json, os, re, tempfile, base64
from pathlib import Path
from docx import Document
from pypdf import PdfReader

MEFORSHIM=['Rashi','Ramban','Sforno','Daas Zekenim','Ohr HaChaim','Ibn Ezra','Chizkuni','Rabbeinu Bachaye']

def read_docx_bytes(data):
    doc=Document(io.BytesIO(data))
    parts=[]
    for p in doc.paragraphs:
        if p.text.strip(): parts.append(p.text.strip())
    for table in doc.tables:
        for row in table.rows:
            txt=' | '.join(c.text.strip() for c in row.cells if c.text.strip())
            if txt: parts.append(txt)
    return '\n'.join(parts)

def read_pdf_text(data):
    try:
        r=PdfReader(io.BytesIO(data))
        return '\n'.join((p.extract_text() or '') for p in r.pages)
    except Exception:
        return ''

def openai_ready():
    return bool(os.getenv('OPENAI_API_KEY')) and bool(os.getenv('OPENAI_MODEL'))

def _client():
    from openai import OpenAI
    return OpenAI(api_key=os.getenv('OPENAI_API_KEY'))

def _clean_json(raw):
    raw=raw.strip(); raw=re.sub(r'^```(?:json)?\s*|\s*```$','',raw,flags=re.S)
    return json.loads(raw)

def extract_with_ai(data, filename, source_type, perek=None):
    if not openai_ready():
        raise RuntimeError('OPENAI_API_KEY and OPENAI_MODEL are required to read scanned/handwritten files.')
    client=_client(); model=os.getenv('OPENAI_MODEL'); suffix=Path(filename).suffix.lower()
    prompt = f'''You are indexing a teacher's Chumash curriculum or test for Sefer Shemot.
File type: {source_type}. Perek hint: {perek or 'unknown'}.
Preserve the teacher's wording and structure. For NOTES, extract the highlighted/main taught ideas and named mefarshim. For TESTS, extract every question and any answer actually present. For scanned student tests, preserve handwritten/typed answers when legible. Never invent missing answers.
Return VALID JSON ONLY:
{{"summary":"brief summary","curriculum_items":[{{"sefer":"Shemot","perek":3,"pasuk_start":1,"pasuk_end":2,"meforash":"Rashi","topic":"topic","details":"details","importance":1-10,"source_type":"notes|sheet"}}],"style_examples":[{{"section":"section name","prompt":"question text","answer":"answer if present","perek":3,"difficulty":1-10}}]}}
'''
    if suffix in ('.pdf', '.png', '.jpg', '.jpeg', '.webp', '.heic'):
        file_suffix = suffix if suffix != '.jpeg' else '.jpg'
        with tempfile.NamedTemporaryFile(suffix=file_suffix,delete=False) as tmp:
            tmp.write(data); tmp.flush(); tmp_path=tmp.name
        try:
            with open(tmp_path,'rb') as f:
                uploaded=client.files.create(file=f,purpose='user_data')
            resp=client.responses.create(
                model=model,
                instructions=prompt,
                input=[{'role':'user','content':[{'type':'input_file','file_id':uploaded.id},{'type':'input_text','text':'Read the entire file visually. It may be a scan or photo. Capture Hebrew, English, handwriting, question numbering, answer lines, mefarshim, and perek/pasuk markers. Do not guess illegible text.'}]}],
            )
            return _clean_json(resp.output_text)
        finally:
            try: os.unlink(tmp_path)
            except Exception: pass
    text=read_docx_bytes(data) if suffix=='.docx' else data.decode('utf-8','ignore')
    resp=client.responses.create(model=model,instructions=prompt,input=text[:60000])
    return _clean_json(resp.output_text)

def ingest_file(data, filename, source_type, perek=None):
    suffix=Path(filename).suffix.lower(); raw=''
    if suffix=='.docx': raw=read_docx_bytes(data)
    elif suffix=='.pdf': raw=read_pdf_text(data)
    elif suffix in ('.txt','.md'): raw=data.decode('utf-8','ignore')
    if openai_ready(): structured=extract_with_ai(data, filename, source_type, perek)
    else: structured={'summary':'Stored without AI indexing','curriculum_items':[],'style_examples':[]}
    return raw, structured
